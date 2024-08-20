import pandas as pd
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE

def add_hyperlink(paragraph, text: str, url: str) -> OxmlElement:
    """
    向段落中添加超链接。

    这个函数向给定的段落对象中添加一个包含指定文本和URL的超链接。它处理了创建XML元素、设置关系ID以及应用超链接样式等细节。

    参数:
        paragraph: 要添加超链接的段落对象 (docx.text.Paragraph)。
        text: 超链接显示的文本。
        url: 超链接的目标URL。

    返回:
        创建的超链接XML元素 (docx.oxml.shared.OxmlElement)。
    """

    # 获取文档的关系文件并生成新的关系ID
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # 创建超链接XML元素并设置关系ID
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # 创建运行元素和属性元素
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 组合XML元素并设置超链接文本
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # 将超链接添加到段落中
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # 设置超链接样式 (如果没有预设样式的话)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def eastmoney_docx_generator(date: str, df: pd.DataFrame) -> None:
    """
    生成包含东方财富新闻信息的 Word 文档。

    这个函数将 pandas DataFrame 中的新闻数据格式化为 Word 文档，包括标题、摘要、公司信息、链接等。

    参数:
        date: 用于生成文件名的日期字符串，格式为 "MM-DD"（例如 "07-30"）。
        df: 包含新闻数据的 pandas DataFrame，应包含以下列：
            - "title": 新闻标题
            - "abstract": 新闻摘要
            - "name": 公司名称
            - "code": 公司代码
            - "web": 新闻链接

    返回:
        None。该函数直接将文档保存到文件。
    """

    filename = f"report/eastmoney/{date} Report.docx"
    doc = Document()

    # 设置默认字体为微软雅黑
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    rFonts = font._element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for _, row in df.iterrows():  # 遍历 DataFrame 的每一行
        # 添加标题，设置样式
        heading = doc.add_heading(level=1)
        for char in row['title'].strip():
            run = heading.add_run(char)
            run.font.name = '微软雅黑'
            run.font.size = Pt(16)
            run.bold = True  # 加粗
            run.font.color.rgb = RGBColor(70, 130, 180)  # 设置颜色
            if ord(char) > 127:  # 判断是否是中文字符
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 设置中文字体

        # 添加摘要
        p = doc.add_paragraph(f"摘要: {row['abstract']}")

        # 添加公司和代码信息
        p.add_run(f"\n公司: {row['name']}    |")
        p.add_run(f"    代码: {row['code']}\n")

        # 添加超链接
        add_hyperlink(p, row['web'], row['web'])

        # 添加空行
        doc.add_paragraph("\n")

    doc.save(filename)
    print(f"Finished Creating {filename}...")


def dongmi_docx_generator(date: str, df: pd.DataFrame) -> None:
    """
    生成包含东方财富董秘问答信息的 Word 文档。

    这个函数将 pandas DataFrame 中的新闻数据格式化为 Word 文档，包括问题、回答、公司名等。

    参数:
        date: 用于生成文件名的日期字符串，格式为 "MM-DD"（例如 "07-30"）。
        df: 包含新闻数据的 pandas DataFrame，应包含以下列：
            - "name": 公司名称
            - "question": 投资者问题
            - "answer": 董秘回答 

    返回:
        None。该函数直接将文档保存到文件。
    """
    filename = f'report/dongmi/{date} 董秘问答.docx'
    doc = Document()

    # 设置默认字体为微软雅黑
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    rFonts = font._element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 遍历Dataframe的每一行
    for _, row in df.iterrows():
        # 添加标题，设置样式
        heading = doc.add_heading(level=1)
        for char in f"公司：{row['name']}".strip():
            run = heading.add_run(char)
            run.font.name = '微软雅黑'
            run.font.size = Pt(16)
            run.bold = True  # 加粗
            run.font.color.rgb = RGBColor(70, 130, 180)  # 设置颜色
            if ord(char) > 127:  # 判断是否是中文字符
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 设置中文字体
        # Add the `question` column as a heading
        doc.add_paragraph(f"问题: {row['question']}")
        # Add the `answer` column as a paragraph
        doc.add_paragraph(f"回答: {row['answer']}")

    doc.save(filename)
    print(f"Finished Creating {filename}...")